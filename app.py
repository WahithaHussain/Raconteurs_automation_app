from datetime import datetime, timedelta
import io
import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
import streamlit as st

# --- STREAMLIT PAGE CONFIG ---
st.set_page_config(
    page_title="BOSCH Raconteurs Toastmasters Agenda & Intros Generator", page_icon="🎤", layout="centered"
)

st.title("🎤 Toastmasters Document Generator")
st.write(
    "Upload the meeting plan to generate both the Agenda and Member Intros Word"
    " documents."
)


# --- HELPER FUNCTIONS (AGENDA GENERATION) ---
def clean_name(raw_text):
  """Removes department codes in parentheses like (MS/ENE-ESH4-XC) and trailing spaces."""
  if not raw_text:
    return ""
  cleaned = re.sub(r"\([^)]*\)", "", raw_text)
  return cleaned.strip()


def extract_speaker_info(raw_text):
  """Extracts clean speaker name, numerical duration for time math, and display duration string.

  Defaults to 7 minutes for time math and "5-7 mins" for display if omitted.
  """
  if not raw_text:
    return "", 7, "5-7 mins"

  time_match = re.search(
      r"(\d+)(?:\s*-\s*(\d+))?\s*mins?", raw_text, re.IGNORECASE
  )
  duration = 7
  dur_display = "5-7 mins"

  if time_match:
    num1 = int(time_match.group(1))
    if time_match.group(2):
      num2 = int(time_match.group(2))
      duration = max(num1, num2)
      dur_display = f"{num1}-{num2} mins"
    else:
      duration = num1
      dur_display = f"{num1} mins"

    raw_text = re.sub(
        r"(\d+)(?:\s*-\s*(\d+))?\s*mins?", "", raw_text, flags=re.IGNORECASE
    )

  return clean_name(raw_text), duration, dur_display


def parse_role_file(txt_content):
  """Parses meeting metadata, standard role players, and dynamic N speakers/evaluators."""
  lines = [
      line.strip()
      for line in txt_content.decode("utf-8").splitlines()
      if line.strip()
  ]

  data = {
      "theme": "",
      "meeting_num": "0",
      "date_str": "",
      "date_code": "",
      "roles": {},
      "speaker_durations": {},
      "speaker_display_durations": {},
      "speaker_indices": [],
  }

  full_text = "\n".join(lines)

  # Extract Theme
  theme_match = re.search(r"THEME\s*:\s*(.*)", full_text, re.IGNORECASE)
  if theme_match:
    data["theme"] = theme_match.group(1).strip()

  # Extract Meeting Number
  num_match = re.search(r"Meeting:\s*#?(\d+)", full_text, re.IGNORECASE)
  if num_match:
    data["meeting_num"] = num_match.group(1)

  # Extract Date
  date_match = re.search(r"(\d{2}\.\d{2}\.\d{4})", full_text)
  if date_match:
    raw_date = date_match.group(1)
    dt = datetime.strptime(raw_date, "%d.%m.%Y")
    data["date_str"] = dt.strftime("%B %d, %Y")
    data["date_code"] = dt.strftime("%d%m%y")

  # Standard non-speaker roles
  std_roles = {
      "SAA": r"SAA\s*:\s*(.*)",
      "PO": r"PO\s*:\s*(.*)",
      "TMOD": r"TMOD\s*:\s*(.*)",
      "GE": r"GE\s*:\s*(.*)",
      "TTM": r"TTM\s*:\s*(.*)",
      "Timer": r"TIMER\s*:\s*(.*)",
      "Ah Counter": r"AH-COUNTER\s*:\s*(.*)",
      "Grammarian": r"GRAMMARIAN\s*:\s*(.*)",
      "Listener": r"Listener\s*:\s*(.*)",
  }

  for role_key, pattern in std_roles.items():
    match = re.search(pattern, full_text, re.IGNORECASE)
    data["roles"][role_key] = clean_name(match.group(1)) if match else ""

  # Dynamic Speakers & Evaluators discovery (Speaker 1, 2, 3, ...)
  speaker_matches = re.findall(
      r"Speaker\s*(\d+)\s*:\s*(.*)", full_text, re.IGNORECASE
  )
  spk_nums = set()

  for num_str, raw_val in speaker_matches:
    num = int(num_str)
    name, duration, dur_display = extract_speaker_info(raw_val)
    if name:  # Only add active speakers
      spk_nums.add(num)
      data["roles"][f"Speaker {num}"] = name
      data["speaker_durations"][f"Speaker {num}"] = duration
      data["speaker_display_durations"][f"Speaker {num}"] = dur_display

      eval_match = re.search(
          rf"Evaluator\s*{num}\s*:\s*(.*)", full_text, re.IGNORECASE
      )
      data["roles"][f"Evaluator {num}"] = (
          clean_name(eval_match.group(1)) if eval_match else ""
      )

  data["speaker_indices"] = sorted(list(spk_nums))
  return data


def build_agenda_schedule(parsed_data):
  """Dynamically constructs agenda slots with custom speech timing and optional TTM/Listener sessions."""
  roles = parsed_data["roles"]
  durations = parsed_data["speaker_durations"]
  display_durations = parsed_data.get("speaker_display_durations", {})
  speaker_indices = parsed_data.get("speaker_indices", [1, 2])

  has_listener = bool(roles.get("Listener"))
  has_ttm = bool(roles.get("TTM"))
  tmod = roles.get("TMOD", "")

  master_items = [
      (15, "15 mins", "Networking & Technology setup", "ALL"),
      (5, "05 mins", "SAA address", roles.get("SAA", "")),
      (
          10,
          "10 mins",
          "Presidential address TMOD introduction",
          f"{roles.get('PO', '')} / {tmod}".strip(" /"),
      ),
      (
          7,
          "07 mins",
          "Meeting structure, Theme address and GE introduction",
          tmod,
      ),
      (2, "02 mins", "GE address and TAG team Intro", roles.get("GE", "")),
      (2, "02 mins", "Timer", roles.get("Timer", "")),
      (2, "02 mins", "Ah Counter", roles.get("Ah Counter", "")),
      (2, "02 mins", "Grammarian", roles.get("Grammarian", "")),
  ]

  if has_listener:
    master_items.append((2, "02 mins", "Listener", roles.get("Listener", "")))

  # Add Prepared Speeches dynamically with timing formatted
  for i in speaker_indices:
    spk_key = f"Speaker {i}"
    eval_key = f"Evaluator {i}"
    spk_name = roles.get(spk_key, "")
    eval_name = roles.get(eval_key, "")
    spk_dur = durations.get(spk_key, 7)
    spk_dur_disp = display_durations.get(spk_key, "5-7 mins")

    master_items.append((
        2,
        "02 mins",
        f"Evaluator {i} intro – Speaker {i} objective – Speaker Intro",
        f"{tmod} – {eval_name}".strip(" –"),
    ))
    master_items.append((
        spk_dur,
        spk_dur_disp,
        f"Prepared Speech {i} ({spk_dur_disp})",
        spk_name,
    ))

  master_items.append((5, "05 mins", "TMOD interaction", tmod))

  # Conditional Table Topics Session
  if has_ttm:
    master_items.append(
        (15, "15 mins", "Table Topics Session", roles.get("TTM", ""))
    )

  for i in speaker_indices:
    eval_name = roles.get(f"Evaluator {i}", "")
    master_items.append((4, "04 mins", f"Evaluator {i} Report", eval_name))

  master_items.append((
      6,
      "06 mins",
      "TAG Team Report",
      (
          f"{roles.get('Timer', '')} / {roles.get('Ah Counter', '')} /"
          f" {roles.get('Grammarian', '')}"
      ),
  ))

  if has_listener:
    master_items.append(
        (5, "05 mins", "Listening Quiz Session", roles.get("Listener", ""))
    )

  master_items.append((5, "05 mins", "GE Report", roles.get("GE", "")))
  master_items.append((2, "02 mins", "TMOD Conclusion", tmod))
  master_items.append((
      8,
      "08 mins",
      "Voting – Meeting closure – Photo session",
      roles.get("PO", ""),
  ))

  active_schedule = []
  current_time = datetime.strptime("17:00", "%H:%M")

  for duration_mins, dur_str, role_text, role_taker in master_items:
    time_str = f"{current_time.hour % 12 or 12}:{current_time.minute:02d}"
    active_schedule.append({
        "time": time_str,
        "duration": dur_str,
        "role_text": role_text,
        "role_taker": role_taker,
    })
    current_time += timedelta(minutes=duration_mins)

  return active_schedule


def set_paragraph_text_and_style(
    paragraph, text, font_size_pt=17, color_rgb=(255, 255, 255), bold=True
):
  paragraph.text = text
  for run in paragraph.runs:
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = bold


# --- HELPER FUNCTIONS (INTRO GENERATION) ---
def load_member_database(excel_source):
  """Loads and standardizes Excel database across all sheets."""
  try:
    xls = pd.ExcelFile(excel_source)
    records = []
    for sheet_name in xls.sheet_names:
      df = pd.read_excel(excel_source, sheet_name=sheet_name)
      name_col, intro_col = None, None
      for col in df.columns:
        for val in df[col].dropna():
          val_str = str(val).strip().lower()
          if val_str == "name":
            name_col = col
          elif val_str == "introduction":
            intro_col = col

      if name_col is not None and intro_col is not None:
        header_idx = df[
            df[name_col].astype(str).str.strip().str.lower() == "name"
        ].index[0]
        clean_df = df.iloc[header_idx + 1 :].copy()
        clean_df = clean_df.rename(
            columns={name_col: "Name", intro_col: "Introduction"}
        )
        records.append(clean_df[["Name", "Introduction"]])

    if records:
      return pd.concat(records, ignore_index=True).dropna(subset=["Name"])
  except Exception:
    pass
  return pd.DataFrame(columns=["Name", "Introduction"])


def match_intro(name, db_df):
  """Matches role taker names to database bios."""
  if not name or name.strip().upper() in ["ALL", ""]:
    return "N/A"

  clean_name = name.strip().lower()

  for _, row in db_df.iterrows():
    db_name = str(row["Name"]).strip().lower()
    if db_name and (clean_name in db_name or db_name in clean_name):
      intro = str(row["Introduction"]).strip()
      if intro and intro.lower() != "nan":
        return intro

  tokens = [t for t in clean_name.split() if len(t) > 2]
  for _, row in db_df.iterrows():
    db_name = str(row["Name"]).strip().lower()
    if any(t in db_name for t in tokens):
      intro = str(row["Introduction"]).strip()
      if intro and intro.lower() != "nan":
        return intro

  return "Introduction not provided in database."


# --- USER INTERFACE & MAIN LOGIC ---
st.subheader("1. Upload Inputs")

# --- INSTRUCTIONS EXPANDER ---
with st.expander(
    "📖 How to format the Role Players .txt file? (Click to view guidance)"
):
  st.markdown("""
    **Quick Rules:**
    * **File Type:** Plain text file (`.txt`) created in Notepad, TextEdit, or VS Code.
    * **Teams Format:** Copy the Teams channel message of role takers updated by VPE into the `.txt` file.
    * **Key-Value Format:** Every line must follow the pattern `KEY : Value`.
    * **Date Format:** Use `DD.MM.YYYY` (e.g., `24.07.2026`).
    * **Optional Roles:** If there is no Speaker 2, Listener, or TTM, leave the line blank or omit it entirely.
    * **Unfilled Roles:** For other roles with no role taker assigned yet, use `XXX` as a placeholder.
    """)

# --- FILE UPLOADERS ---
txt_file = st.file_uploader(
    "Upload Role Players Plan (.txt)", type=["txt"], help="Mandatory file"
)
excel_file = st.file_uploader(
    "Upload Member Database (.xlsx) - OPTIONAL",
    type=["xlsx"],
    help="Optional. Fallback server database will be used if skipped.",
)

if txt_file is not None:
  # Determine Excel database source
  if excel_file is not None:
    db_df = load_member_database(excel_file)
    st.info("ℹ️ Using newly uploaded Excel database.")
  elif os.path.exists("Toastmasters Introduction .xlsx"):
    db_df = load_member_database("Toastmasters Introduction .xlsx")
    st.info("ℹ️ Using default server Excel database.")
  else:
    db_df = pd.DataFrame(columns=["Name", "Introduction"])
    st.warning("⚠️ No Excel database found. Placeholders will be used.")

  parsed = parse_role_file(txt_file.getvalue())
  st.success(
      f"Parsed Meeting #{parsed['meeting_num']} | Theme: {parsed['theme']}"
  )

  # --- GENERATE AGENDA DOC ---
  template_path = os.path.join(
      "template", "ToastMastersAgenda_template.docx"
  )
  if os.path.exists(template_path):
    doc = Document(template_path)

    set_paragraph_text_and_style(
        doc.paragraphs[0],
        f"Bosch Raconteurs Toastmasters Club Meeting #{parsed['meeting_num']}",
        font_size_pt=17,
        color_rgb=(255, 255, 255),
        bold=True,
    )
    set_paragraph_text_and_style(
        doc.paragraphs[1],
        f"THEME: {parsed['theme']}",
        font_size_pt=17,
        color_rgb=(255, 255, 255),
        bold=True,
    )
    doc.paragraphs[4].text = parsed["date_str"]

    active_schedule = build_agenda_schedule(parsed)
    table = doc.tables[0]

    # Ensure table has enough rows for active schedule
    while len(table.rows) - 1 < len(active_schedule):
      table.add_row()

    for idx, item in enumerate(active_schedule, start=1):
      row_cells = table.rows[idx].cells
      row_cells[0].text = item["time"]
      row_cells[1].text = item["duration"]
      row_cells[2].text = item["role_text"]
      row_cells[3].text = item["role_taker"]

    total_needed_rows = len(active_schedule) + 1
    while len(table.rows) > total_needed_rows:
      extra_row = table.rows[-1]
      extra_row._tr.getparent().remove(extra_row._tr)

    agenda_buffer = io.BytesIO()
    doc.save(agenda_buffer)
    agenda_buffer.seek(0)

    # --- GENERATE INTROS DOC ---
    intro_doc = Document()
    for s in intro_doc.sections:
      s.top_margin = Inches(1)
      s.bottom_margin = Inches(1)
      s.left_margin = Inches(1)
      s.right_margin = Inches(1)

    p_head = intro_doc.add_paragraph()
    r_head = p_head.add_run(
        f"Bosch Raconteurs Toastmasters meeting #{parsed['meeting_num']} member"
        " intros"
    )
    r_head.font.name = "Calibri"
    r_head.font.size = Pt(20)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(31, 78, 121)
    p_head.paragraph_format.space_after = Pt(18)

    role_map = parsed["roles"]

    # Build TMOD roles dynamically based on parsed speakers
    tmod_roles = []
    for i in parsed.get("speaker_indices", [1, 2]):
      if role_map.get(f"Speaker {i}"):
        tmod_roles.append((f"Speaker {i}", role_map[f"Speaker {i}"]))
      if role_map.get(f"Evaluator {i}"):
        tmod_roles.append((f"Evaluator {i}", role_map[f"Evaluator {i}"]))

    if role_map.get("TTM"):
      tmod_roles.append(("Table Topics Master (TTM)", role_map["TTM"]))

    tmod_roles.append(("General Evaluator (GE)", role_map.get("GE", "N/A")))

    # GE roles (conditional Listener)
    ge_roles = [
        ("Timer", role_map.get("Timer", "N/A")),
        ("Ah-Counter", role_map.get("Ah Counter", "N/A")),
        ("Grammarian", role_map.get("Grammarian", "N/A")),
    ]
    if role_map.get("Listener"):
      ge_roles.append(("Listener", role_map["Listener"]))

    sections_schema = [
        (
            "1. Introduction of PO",
            [("Presiding Officer (PO)", role_map.get("PO", "N/A"))],
        ),
        (
            "2. Introduction of TMOD",
            [("Toastmaster of the Day (TMOD)", role_map.get("TMOD", "N/A"))],
        ),
        ("3. Roles Introduced by TMOD", tmod_roles),
        ("4. Roles Introduced by GE", ge_roles),
    ]

    for cat_title, role_tuples in sections_schema:
      cp = intro_doc.add_paragraph()
      c_run = cp.add_run(cat_title)
      c_run.font.name = "Calibri"
      c_run.font.size = Pt(14)
      c_run.font.bold = True
      c_run.font.color.rgb = RGBColor(46, 117, 182)
      cp.paragraph_format.space_before = Pt(14)
      cp.paragraph_format.space_after = Pt(6)

      for role_label_text, taker_name in role_tuples:
        intro_text = match_intro(taker_name, db_df)

        rp = intro_doc.add_paragraph()
        rp.paragraph_format.left_indent = Inches(0.2)
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after = Pt(2)

        r_lbl = rp.add_run(f"• {role_label_text}: ")
        r_lbl.font.name = "Calibri"
        r_lbl.font.size = Pt(11)
        r_lbl.font.bold = True

        r_name = rp.add_run(taker_name)
        r_name.font.name = "Calibri"
        r_name.font.size = Pt(11)
        r_name.font.bold = True
        r_name.font.color.rgb = RGBColor(31, 78, 121)

        ip = intro_doc.add_paragraph()
        ip.paragraph_format.left_indent = Inches(0.4)
        ip.paragraph_format.space_after = Pt(8)

        irun = ip.add_run(intro_text)
        irun.font.name = "Calibri"
        irun.font.size = Pt(10.5)

        if "not provided" in intro_text.lower() or intro_text == "N/A":
          irun.font.bold = True
          irun.font.italic = False
          irun.font.color.rgb = RGBColor(192, 0, 0)
        else:
          irun.font.bold = False
          irun.font.italic = True
          irun.font.color.rgb = RGBColor(89, 89, 89)

    intro_buffer = io.BytesIO()
    intro_doc.save(intro_buffer)
    intro_buffer.seek(0)

    # --- DOWNLOAD BUTTONS ---
    st.subheader("2. Download Generated Documents")
    col1, col2 = st.columns(2)
    with col1:
      st.download_button(
          label="📄 Download Meeting Agenda",
          data=agenda_buffer,
          file_name=(
              f"ToastMastersAgenda_{parsed['meeting_num']}_{parsed['date_code']}.docx"
          ),
          mime=(
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          ),
      )
    with col2:
      st.download_button(
          label="📝 Download Member Intros",
          data=intro_buffer,
          file_name=(
              f"Bosch_Raconteurs_Meeting_{parsed['meeting_num']}_Member_Intros.docx"
          ),
          mime=(
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          ),
      )
  else:
    st.error(
        "Template missing on server! Ensure"
        " 'template/ToastMastersAgenda_template.docx' exists."
    )
